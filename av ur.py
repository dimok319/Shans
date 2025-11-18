import pandas as pd
from docx import Document
import os
from datetime import datetime
import math

# --- Конфигурация путей ---
base_path = r"\\192.168.1.211\аналитический центр\Отчёты\1 Проекты\Иски для юр.отдела"

# Правильные пути
template_path = os.path.join(base_path, "иск поставка.docx")
data_file_path = os.path.join(base_path, "Обработанный_итоговый_файл.xlsx")
output_folder = os.path.join(base_path, "готовые_заявления")

# Создаем папку для готовых файлов, если ее нет
os.makedirs(output_folder, exist_ok=True)


def format_currency(amount):
    """Форматирует число в строку с разделителями разрядов и добавляет 'рублей'/'копеек'."""
    try:
        rub = int(amount)
        kop = round((amount - rub) * 100)
        rub_str = f"{rub:,}".replace(',', ' ')
        currency_str = f"{rub_str} рублей {kop:02d} коп."
        return currency_str
    except (ValueError, TypeError):
        return "0 рублей 00 коп."


def num2words(num):
    """Преобразует число в слова на русском языке (упрощенная версия)."""
    try:
        num = int(num)
    except (ValueError, TypeError):
        return "ноль"

    ones = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
    tens = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
    teens = ['десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать', 'шестнадцать',
             'семнадцать', 'восемнадцать', 'девятнадцать']
    hundreds = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот', 'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']
    thousands = ['', 'тысяча', 'тысячи', 'тысяч']
    millions = ['', 'миллион', 'миллиона', 'миллионов']

    def triple_to_words(n, is_thousands=False):
        if n == 0:
            return ''
        words = []
        # Сотни
        if n // 100 > 0:
            words.append(hundreds[n // 100])
        n %= 100
        # Десятки и единицы
        if 10 <= n < 20:
            words.append(teens[n - 10])
        else:
            if n // 10 > 0:
                words.append(tens[n // 10])
            if n % 10 > 0:
                word = ones[n % 10]
                if is_thousands:
                    if n % 10 == 1:
                        word = 'одна'
                    elif n % 10 == 2:
                        word = 'две'
                words.append(word)
        return ' '.join(words).strip()

    if num == 0:
        return 'ноль'

    # Разбиваем число на миллионы, тысячи и единицы
    mill = num // 1_000_000
    thous = (num % 1_000_000) // 1_000
    units = num % 1_000

    parts = []

    # Миллионы
    if mill > 0:
        part = triple_to_words(mill)
        if mill % 10 == 1 and mill % 100 != 11:
            part += f' {millions[1]}'
        elif 2 <= mill % 10 <= 4 and (mill % 100 < 10 or mill % 100 >= 20):
            part += f' {millions[2]}'
        else:
            part += f' {millions[3]}'
        parts.append(part)

    # Тысячи
    if thous > 0:
        part = triple_to_words(thous, is_thousands=True)
        if thous % 10 == 1 and thous % 100 != 11:
            part += f' {thousands[1]}'
        elif 2 <= thous % 10 <= 4 and (thous % 100 < 10 or thous % 100 >= 20):
            part += f' {thousands[2]}'
        else:
            part += f' {thousands[3]}'
        parts.append(part)

    # Единицы
    if units > 0 or (mill == 0 and thous == 0):
        parts.append(triple_to_words(units))

    return ' '.join(parts).strip()


# --- Проверка существования файлов ---
print("=" * 50)
print("ПРОВЕРКА ФАЙЛОВ:")
print(f"Шаблон: {template_path}")
print(f"Существует: {os.path.exists(template_path)}")
print(f"Данные: {data_file_path}")
print(f"Существует: {os.path.exists(data_file_path)}")
print(f"Папка для результатов: {output_folder}")
print("=" * 50)

if not os.path.exists(template_path):
    print(f"❌ ОШИБКА: Файл шаблона не найден: {template_path}")
    exit()

if not os.path.exists(data_file_path):
    print(f"❌ ОШИБКА: Файл с данными не найден: {data_file_path}")
    exit()

# --- Чтение данных из Excel ---
try:
    # Сначала узнаем названия всех листов
    excel_file = pd.ExcelFile(data_file_path)
    print(f"📊 Доступные листы в файле: {excel_file.sheet_names}")

    # Читаем лист 'Сводный отчет' как вы указали
    if 'Сводный отчет' in excel_file.sheet_names:
        df = pd.read_excel(data_file_path, sheet_name='Сводный отчет')
        print("✅ Читаем лист 'Сводный отчет'")
    else:
        print("❌ Лист 'Сводный отчет' не найден")
        print("📋 Доступные листы:", excel_file.sheet_names)
        exit()

except Exception as e:
    print(f"❌ Ошибка при чтении Excel файла: {e}")
    exit()

# --- Проверяем наличие необходимых столбцов ---
print("🔍 Проверяем столбцы в таблице...")
print(f"📋 Существующие столбцы: {list(df.columns)}")

required_columns = ['Контрагент', 'Истец_Наименование', 'Истец_Адрес', 'Истец_ИНН',
                    'Юр адрес', 'ИНН', 'Дата договора', 'Номер', 'Директор_Имя',
                    'Дата_Расчета', 'Период_Начало', 'Госпошлина', 'Приложения',
                    'Остаток задолженности', 'Сумма отгружено']

missing_columns = [col for col in required_columns if col not in df.columns]
if missing_columns:
    print(f"❌ ОШИБКА: В таблице отсутствуют необходимые столбцы: {missing_columns}")
    exit()
else:
    print("✅ Все необходимые столбцы найдены")

# --- Поиск данных для ООО «НИВА» ---
target_company = 'ООО «НИВА»'
print(f"🔎 Ищем данные для: {target_company}")

company_data = df[df['Контрагент'] == target_company]

if company_data.empty:
    print(f"❌ Данные для контрагента '{target_company}' не найдены.")
    print(f"📋 Доступные контрагенты: {df['Контрагент'].unique()}")
    exit()

print(f"✅ Найдено записей: {len(company_data)}")

# Берем первую найденную запись
row = company_data.iloc[0]

# --- Заполнение данных ---
print("🔄 Заполняем данные в шаблон...")

# Данные из строки таблицы
data_map = {
    '{istec_name}': str(row['Истец_Наименование']),
    '{istec_address}': str(row['Истец_Адрес']),
    '{istec_inn}': str(row['Истец_ИНН']),
    '{otvetchik_name}': str(row['Контрагент']),
    '{otvetchik_address}': str(row['Юр адрес']),
    '{otvetchik_inn}': str(row['ИНН']),
    '{dogovor_date}': str(row['Дата договора']),
    '{dogovor_num}': str(row['Номер']),
    '{director_name}': str(row['Директор_Имя']),
    '{data_rascheta}': str(row['Дата_Расчета']),
    '{period_start}': str(row['Период_Начало']),
    '{gosposhlina_rub}': format_currency(row['Госпошлина']),
    '{primen_list}': str(row['Приложения']),
}

# Рассчитанные данные (основной долг, проценты, итого)
osnovnoy_dolg = row['Остаток задолженности']
data_map['{osnovnoy_dolg_rub}'] = format_currency(osnovnoy_dolg)
data_map['{osnovnoy_dolg_words}'] = num2words(int(osnovnoy_dolg))

# --- ЗДЕСЬ ДОЛЖНА БЫТЬ ВАША ЛОГИКА РАСЧЕТА ПРОЦЕНТОВ ---
# Она сложная, поэтому для примера возьмем упрощенный вариант
procenty_sum = 1399325.53  # Рассчитайте эту сумму по вашим правилам
data_map['{procenty_sum_rub}'] = format_currency(procenty_sum)
data_map['{procenty_sum_words}'] = num2words(int(procenty_sum))

obshaya_zadolzhennost = osnovnoy_dolg + procenty_sum
data_map['{obshaya_zadolzhennost_rub}'] = format_currency(obshaya_zadolzhennost)
data_map['{obshaya_zadolzhennost_words}'] = num2words(int(obshaya_zadolzhennost))

data_map['{tsena_iska_rub}'] = data_map['{obshaya_zadolzhennost_rub}']
data_map['{summa_otgruzeno_rub}'] = format_currency(row['Сумма отгружено'])
data_map['{summa_otgruzeno_words}'] = num2words(int(row['Сумма отгружено']))

# --- Заполнение шаблона Word ---
try:
    print("📝 Открываем шаблон Word...")
    doc = Document(template_path)

    print("🔄 Заменяем текст в параграфах...")
    for paragraph in doc.paragraphs:
        for key, value in data_map.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    print("🔄 Заменяем текст в таблицах...")
    for table in doc.tables:
        for row_table in table.rows:
            for cell in row_table.cells:
                for key, value in data_map.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # --- Сохранение результата ---
    today_date = datetime.now().strftime("%d.%m.%Y")
    output_filename = f"Исковое заявление {target_company} от {today_date}.docx"
    output_path = os.path.join(output_folder, output_filename)

    print("💾 Сохраняем файл...")
    doc.save(output_path)
    print("=" * 50)
    print(f"✅ Исковое заявление успешно сохранено!")
    print(f"📁 Файл: {output_path}")
    print("=" * 50)

except Exception as e:
    print(f"❌ Ошибка при работе с Word документом: {e}")