import os
import re
from decimal import Decimal, ROUND_HALF_UP

import pandas as pd
from docx import Document
from num2words import num2words


# =========================
# НАСТРОЙКИ ВХОДА/ВЫХОДА
# =========================
EXCEL_FILE = r'C:\Users\nkazakov\Downloads\Отчет по ДЗ 06.10.2025.xlsx'
SHEET_NAME = 'Реестр ПДЗ'
TEMPLATE_PATH = r'C:\Users\nkazakov\Downloads\Претензия НОВАЯ.docx'
OUTPUT_FOLDER = r'\\192.168.1.211\файлообменный ресурс\Внутренний обмен\Юр_документы'
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


# =========================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =========================
def sanitize_filename(filename: str, max_length: int = 100) -> str:
    safe = re.sub(r'[<>:"/\\|?*]', '_', str(filename))
    return safe[:max_length] if len(safe) > max_length else safe


def ru_plural(n: int, forms: tuple[str, str, str]) -> str:
    """
    Подбор правильной формы слова по числу (1, 2-4, 5+).
    forms = ('рубль','рубля','рублей') или ('копейка','копейки','копеек')
    """
    n_abs = abs(n)
    if 11 <= (n_abs % 100) <= 14:
        return forms[2]
    last = n_abs % 10
    if last == 1:
        return forms[0]
    if 2 <= last <= 4:
        return forms[1]
    return forms[2]


def amount_parts(amount) -> dict:
    """
    Возвращает части суммы:
      - rubles: целые рубли (int)
      - kopeks: копейки (00..99)
      - number_text: '12 345'
      - words_text: 'Двенадцать тысяч триста сорок пять'
      - combined_basic: '12 345 (Двенадцать тысяч триста сорок пять)'
      - combined_full:  '12 345 (Двенадцать тысяч триста сорок пять) рублей 00 копеек'
                         (с правильными падежами для рублей/копеек)
    """
    dec = Decimal(str(amount)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    rubles = int(dec // 1)
    kopeks = int((dec - rubles) * 100)

    number_text = f"{rubles:,.0f}".replace(",", " ")
    words_text = num2words(rubles, lang='ru').capitalize()

    rub_word = ru_plural(rubles, ('рубль', 'рубля', 'рублей'))
    kop_word = ru_plural(kopeks, ('копейка', 'копейки', 'копеек'))

    combined_basic = f"{number_text} ({words_text})"
    combined_full = f"{combined_basic} {rub_word} {kopeks:02d} {kop_word}"

    return {
        'rubles': rubles,
        'kopeks': kopeks,
        'number_text': number_text,
        'words_text': words_text,
        'combined_basic': combined_basic,
        'combined_full': combined_full,
    }


def iter_all_paragraphs(doc: Document):
    """Итерирует абзацы в документе, таблицах, шапках и подвалах."""
    # Тело
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p
    # Шапки/подвалы всех секций
    for s in doc.sections:
        hdr = s.header
        if hdr:
            for p in hdr.paragraphs:
                yield p
            for t in hdr.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            yield p
        ftr = s.footer
        if ftr:
            for p in ftr.paragraphs:
                yield p
            for t in ftr.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            yield p


def replace_placeholders_loose(text: str, mapping: dict) -> str:
    """
    Заменяет {{КЛЮЧ}} на значение, даже если внутри фигурных скобок есть пробелы или NBSP.
    Пример: '{{  СуммаПДЗ  }}', '{{\u00A0СуммаПДЗ\u00A0}}'
    """
    for key, val in mapping.items():
        base = key.strip("{}")  # из '{{СуммаПДЗ}}' -> 'СуммаПДЗ'
        pattern = r"\{\{(?:\s|\u00A0)*" + re.escape(base) + r"(?:\s|\u00A0)*\}\}"
        text = re.sub(pattern, str(val), text)
    return text


def set_tnr_font(doc: Document):
    """Ставит шрифт Times New Roman для всех run-ов (тело/таблицы только)."""
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = 'Times New Roman'
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
    # В шапках/подвалах сохраняем форматирование как есть (обычно в шаблоне уже TNR)


def fill_word_template(row, template_path: str, save_path: str):
    doc = Document(template_path)

    amt = amount_parts(row['ПДЗ, руб'])
    # Для {{СуммаПДЗ}} используем ПОЛНУЮ форму (с "рублей ХХ копеек"),
    # а для «пропусков» ______(______) используем BASIC без валют.
    amount_basic = amt['combined_basic']
    amount_full = amt['combined_full']

    replacements = {
        "{{Контрагент}}": str(row.get("Контрагент", "")),
        "{{Адрес}}": str(row.get("Адрес", "")),
        "{{ИНН}}": str(row.get("ИНН", "")),
        "{{Договор}}": str(row.get("Договор", "")),
        "{{СуммаПДЗ}}": amount_full,                 # например: "12 345 (...) рублей 00 копеек"
        "{{СуммаПДЗ_числом}}": amt['number_text'],   # "12 345"
        "{{СуммаПДЗ_прописью}}": amt['words_text'],  # "Двенадцать тысяч ..."
    }

    # Проход по всем абзацам
    for p in iter_all_paragraphs(doc):
        new_text = replace_placeholders_loose(p.text, replacements)
        # Поддержка «пропусков» вида _____________(_________)
        # Подставляем BASIC, т.к. после пропуска в шаблоне обычно стоит "рублей 00 копеек"
        new_text = re.sub(r'_{3,}\s*\(\s*_{3,}\s*\)', amount_basic, new_text)
        if new_text != p.text:
            p.text = new_text

    # Шрифт
    set_tnr_font(doc)

    # (Необязательно) Проверка остаточных плейсхолдеров — поможет поймать текстовые блоки/фигуры
    leftovers = []
    for p in iter_all_paragraphs(doc):
        leftovers += re.findall(r"\{\{[^}]+\}\}", p.text)
    if leftovers:
        print(f"Внимание: в файле {os.path.basename(save_path)} остались плейсхолдеры: {set(leftovers)}.\n"
              f"Часто это означает, что текст находится в TextBox/фигуре, которую python-docx не читает. "
              f"Переведите такие блоки в режим «В тексте» в шаблоне.")

    doc.save(save_path)


# =========================
# ОСНОВНОЙ СЦЕНАРИЙ
# =========================
def main():
    # 1) Читаем Excel
    df = pd.read_excel(EXCEL_FILE, sheet_name=SHEET_NAME, skiprows=9)

    # 2) Приводим сумму к числу
    df['ПДЗ, руб'] = pd.to_numeric(df['ПДЗ, руб'], errors='coerce').fillna(0)

    # 3) Фильтруем > 0
    df_filtered = df[df['ПДЗ, руб'] > 0]

    # 4) Группируем
    grouped_df = (
        df_filtered
        .groupby(['Контрагент', 'Договор'], dropna=False)
        .agg({'ПДЗ, руб': 'sum', 'Адрес': 'first', 'ИНН': 'first'})
        .reset_index()
    )

    # 5) Генерация документов
    for _, row in grouped_df.iterrows():
        kontr = row['Контрагент'] if pd.notna(row['Контрагент']) else 'Без_контрагента'
        dog = row['Договор'] if pd.notna(row['Договор']) and str(row['Договор']).strip() else 'Без_договора'

        customer_name = sanitize_filename(str(kontr))
        contract_name = sanitize_filename(str(dog))

        base_filename = f"{customer_name}_{contract_name}_документ"
        save_path = os.path.join(OUTPUT_FOLDER, f"{base_filename}.docx")

        # Если файл уже существует — добавим счётчик
        counter = 1
        while os.path.exists(save_path):
            save_path = os.path.join(OUTPUT_FOLDER, f"{base_filename}_{counter}.docx")
            counter += 1

        try:
            fill_word_template(row, TEMPLATE_PATH, save_path)
            print(f"Создан файл: {save_path}")
        except Exception as e:
            print(f"Ошибка при создании файла для '{customer_name} / {contract_name}': {e}")


if __name__ == "__main__":
    main()
